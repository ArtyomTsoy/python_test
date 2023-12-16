from docx import Document
from docx.shared import Pt

doc = Document('work_file.docx')
bold_text = [run.text for para in doc.paragraphs for run in para.runs if run.bold]
bold_text = ''.join(bold_text)
print(bold_text)

doc = Document()
paragraph = doc.add_paragraph('Это новый текст')
run = paragraph.runs[0]
run.font.size = Pt(35)
run.font.name = 'Bahnschrift'
doc.save('output.docx')